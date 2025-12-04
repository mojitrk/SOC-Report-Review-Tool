#!/usr/bin/env python3
"""
Create a fake SOC 2 report for testing
"""

from docx import Document
from docx.shared import Pt

def create_test_report():
    doc = Document()
    
    # Cover page
    doc.add_heading('SOC 2 Type II Report', level=1)
    doc.add_paragraph()
    doc.add_paragraph('Report on Controls Related to Security, Availability, Processing Integrity, Confidentiality, and Privacy')
    doc.add_paragraph()
    
    doc.add_heading('Acme Cloud Services, Inc.', level=2)
    doc.add_paragraph('Period: January 1, 2023 to December 31, 2023')
    doc.add_paragraph('Report Date: March 15, 2024')
    
    doc.add_page_break()
    
    # Management's Assertion
    doc.add_heading("Management's Assertion", level=1)
    doc.add_paragraph(
        'Acme Cloud Services, Inc. (the Company) asserts that it maintains a system of controls, '
        'consistent with the Trust Services Criteria for Security, Availability, Processing Integrity, '
        'Confidentiality, and Privacy, that are suitable to achieve the Company\'s objective of providing '
        'its customers with products and services in accordance with the commitments in its service organization\'s contracts.'
    )
    
    doc.add_page_break()
    
    # Auditor's Opinion
    doc.add_heading("Auditor's Opinion", level=1)
    doc.add_paragraph(
        'We have examined the accompanying Management Assertion regarding the effectiveness of Acme Cloud Services, Inc.\'s '
        'controls related to Security, Availability, Processing Integrity, Confidentiality, and Privacy during the period '
        'January 1, 2023 to December 31, 2023.'
    )
    doc.add_paragraph(
        'Based on our examination, we assert with reasonable assurance that Acme Cloud Services, Inc. maintained effective controls '
        'during the specified period to achieve the objectives stated in the Management Assertion.'
    )
    
    doc.add_page_break()
    
    # Trust Services Criteria - Security
    doc.add_heading('Trust Services Criteria - Security', level=1)
    doc.add_paragraph(
        'The Company has implemented controls to ensure that system access is restricted to authorized personnel. '
        'Multi-factor authentication (MFA) is required for all administrative access. Access reviews are performed quarterly.'
    )
    
    # Trust Services Criteria - Availability
    doc.add_heading('Trust Services Criteria - Availability', level=1)
    doc.add_paragraph(
        'The Company employs 24/7 monitoring of system infrastructure. Alerts are configured for critical metrics. '
        'The Company maintains a Recovery Time Objective (RTO) of 4 hours and Recovery Point Objective (RPO) of 1 hour.'
    )
    
    # Trust Services Criteria - Processing Integrity
    doc.add_heading('Trust Services Criteria - Processing Integrity', level=1)
    doc.add_paragraph(
        'Input validation is implemented at all application entry points. Automated validation rules check for data type, format, and range. '
        'Failed validations are logged and rejected.'
    )
    
    # Trust Services Criteria - Confidentiality
    doc.add_heading('Trust Services Criteria - Confidentiality', level=1)
    doc.add_paragraph(
        'All data is classified and encrypted. Encryption at rest uses AES-256, and encryption in transit uses TLS 1.2 or higher. '
        'Encryption keys are managed through AWS Key Management Service (KMS).'
    )
    
    # Trust Services Criteria - Privacy
    doc.add_heading('Trust Services Criteria - Privacy', level=1)
    doc.add_paragraph(
        'The Company maintains a comprehensive privacy policy that outlines how personal information is collected, used, retained, and shared. '
        'Customer consent is obtained at the time of data collection.'
    )
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'In conclusion, Acme Cloud Services, Inc. has implemented and maintained effective controls related to Security, Availability, '
        'Processing Integrity, Confidentiality, and Privacy during the period January 1, 2023 to December 31, 2023.'
    )
    
    # Save
    doc.save('test_report.docx')
    print("✓ Created test_report.docx")

if __name__ == '__main__':
    create_test_report()
