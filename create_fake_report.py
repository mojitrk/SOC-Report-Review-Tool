#!/usr/bin/env python3
"""
Generate a realistic 10-page fake SOC 2 report for testing purposes.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import random

def add_heading(doc, text, level=1):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to the document."""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def create_fake_soc2_report():
    """Generate a 10-page fake SOC 2 report."""
    doc = Document()
    
    # Cover page
    title = doc.add_paragraph()
    title_run = title.add_run("SOC 2 Type II Report")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Report on Controls Related to Security, Availability, Processing Integrity, Confidentiality, and Privacy")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    company_name = doc.add_paragraph()
    company_name_run = company_name.add_run("Acme Cloud Services, Inc.")
    company_name_run.font.size = Pt(16)
    company_name_run.font.bold = True
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Audit period
    start_date = datetime(2023, 1, 1)
    end_date = datetime(2023, 12, 31)
    audit_period = doc.add_paragraph()
    audit_period_run = audit_period.add_run(f"Period: January 1, 2023 to December 31, 2023")
    audit_period_run.font.size = Pt(12)
    audit_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    report_date = doc.add_paragraph()
    report_date_run = report_date.add_run(f"Report Date: March 15, 2024")
    report_date_run.font.size = Pt(12)
    report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break
    doc.add_page_break()
    
    # Page 2: Management's Assertion
    add_heading(doc, "Management's Assertion", level=1)
    
    doc.add_paragraph(
        "Acme Cloud Services, Inc. (the Company) asserts that it maintains a system of controls, "
        "consistent with the Trust Services Criteria for Security, Availability, Processing Integrity, "
        "Confidentiality, and Privacy (the Criteria), that are suitable to achieve the Company's "
        "objective of providing its customers with products and services in accordance with the commitments "
        "in its service organization's contracts or other agreements."
    )
    
    doc.add_paragraph(
        "This assertion is based on the criteria set forth in Trust Services Criteria established by the "
        "American Institute of Certified Public Accountants (AICPA) and confirmed by the Trust Services Maturity Model."
    )
    
    doc.add_page_break()
    
    # Page 3: Auditor's Opinion
    add_heading(doc, "Auditor's Opinion", level=1)
    
    doc.add_paragraph(
        "We have examined the accompanying Management Assertion regarding the effectiveness of Acme Cloud Services, Inc.'s "
        "controls related to Security, Availability, Processing Integrity, Confidentiality, and Privacy during the period "
        "January 1, 2023 to December 31, 2023, in accordance with attestation standards established by the AICPA."
    )
    
    doc.add_paragraph(
        "Based on our examination and the evidence we obtained, we assert with reasonable assurance that Acme Cloud Services, Inc. "
        "maintained effective controls during the specified period to achieve the objectives stated in the Management Assertion."
    )
    
    doc.add_paragraph(
        "The controls described in this report address the security, availability, processing integrity, confidentiality, and privacy "
        "of the Company's system and the processing of information on behalf of customers."
    )
    
    doc.add_page_break()
    
    # Page 4: Service Organization Description
    add_heading(doc, "Service Organization and System Description", level=1)
    
    doc.add_paragraph(
        "Acme Cloud Services, Inc. is a cloud-based software-as-a-service (SaaS) provider offering enterprise resource planning (ERP) solutions "
        "to mid-market companies across various industries."
    )
    
    doc.add_heading("System Infrastructure", level=2)
    doc.add_paragraph(
        "The Company operates a multi-tenant cloud infrastructure hosted on Amazon Web Services (AWS) in the US-East and US-West regions. "
        "The system includes the following components:"
    )
    
    components = [
        "Web Application Tier: Running on EC2 instances behind an Elastic Load Balancer",
        "Application Tier: Microservices architecture containerized with Docker and orchestrated via Kubernetes",
        "Database Tier: RDS PostgreSQL with Multi-AZ deployment and automated backups",
        "Storage Tier: S3 buckets for document and file storage with encryption at rest",
        "Network: VPC with strict security groups and network ACLs"
    ]
    
    for component in components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_page_break()
    
    # Page 5: Trust Services Criteria - Security
    add_heading(doc, "Trust Services Criteria - Security", level=1)
    
    doc.add_paragraph(
        "The Company has implemented controls to ensure that system access is restricted to authorized personnel and that "
        "unauthorized access attempts are detected and prevented."
    )
    
    doc.add_heading("CC1: Risk Assessment", level=2)
    doc.add_paragraph(
        "Management performs annual risk assessments to identify and evaluate potential threats and vulnerabilities to the system. "
        "Risk assessment results are documented and reviewed by leadership on a quarterly basis."
    )
    
    doc.add_heading("CC2: Logical Access Control", level=2)
    doc.add_paragraph(
        "The Company enforces role-based access control (RBAC) with principle of least privilege. Multi-factor authentication (MFA) "
        "is required for all administrative access. Access reviews are performed quarterly and access is revoked upon termination."
    )
    
    doc.add_page_break()
    
    # Page 6: Trust Services Criteria - Availability
    add_heading(doc, "Trust Services Criteria - Availability", level=1)
    
    doc.add_paragraph(
        "The Company has implemented controls to provide availability of system resources to process information as committed "
        "and required by stakeholders."
    )
    
    doc.add_heading("A1: System Monitoring", level=2)
    doc.add_paragraph(
        "The Company employs 24/7 monitoring of system infrastructure and application health. Alerts are configured for CPU usage, "
        "memory utilization, disk space, and network connectivity. Critical alerts are escalated immediately to on-call engineers."
    )
    
    doc.add_heading("A2: Disaster Recovery and Business Continuity", level=2)
    doc.add_paragraph(
        "A comprehensive Disaster Recovery Plan is maintained and tested annually. The Company maintains a Recovery Time Objective (RTO) "
        "of 4 hours and Recovery Point Objective (RPO) of 1 hour. Backups are performed continuously and stored in geographically diverse locations."
    )
    
    doc.add_page_break()
    
    # Page 7: Trust Services Criteria - Processing Integrity
    add_heading(doc, "Trust Services Criteria - Processing Integrity", level=1)
    
    doc.add_paragraph(
        "The Company has implemented controls to ensure that system processing is complete, accurate, timely, and authorized."
    )
    
    doc.add_heading("PI1: Data Validation", level=2)
    doc.add_paragraph(
        "Input validation is implemented at all application entry points to ensure data accuracy and completeness. Automated validation rules "
        "check for data type, format, range, and referential integrity. Failed validations are logged and rejected."
    )
    
    doc.add_heading("PI2: Error Handling and Reconciliation", level=2)
    doc.add_paragraph(
        "Errors are captured, logged, and escalated based on severity. Reconciliation processes are automated and run hourly to detect discrepancies. "
        "Monthly management reconciliation reports are reviewed for completeness and accuracy."
    )
    
    doc.add_page_break()
    
    # Page 8: Trust Services Criteria - Confidentiality & Privacy
    add_heading(doc, "Trust Services Criteria - Confidentiality and Privacy", level=1)
    
    doc.add_paragraph(
        "The Company has implemented controls to ensure that personal information and confidential data are protected from unauthorized access and use."
    )
    
    doc.add_heading("C1: Data Classification and Encryption", level=2)
    doc.add_paragraph(
        "All data is classified as Public, Internal, Confidential, or Restricted. Encryption at rest uses AES-256, and encryption in transit uses TLS 1.2 or higher. "
        "Encryption keys are managed through AWS Key Management Service (KMS) with key rotation every 90 days."
    )
    
    doc.add_heading("P1: Privacy Policy and Consent", level=2)
    doc.add_paragraph(
        "The Company maintains a comprehensive privacy policy that outlines how personal information is collected, used, retained, and shared. "
        "Customer consent is obtained at the time of data collection, and individuals have the right to access, correct, or delete their data."
    )
    
    doc.add_page_break()
    
    # Page 9: Control Test Results Summary
    add_heading(doc, "Summary of Control Testing", level=1)
    
    doc.add_paragraph(
        "During our examination period (January 1, 2023 - December 31, 2023), we tested a sample of controls across all Trust Service Criteria. "
        "The following table summarizes the results of our testing:"
    )
    
    # Add a simple table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Control Objective'
    header_cells[1].text = 'Tests Performed'
    header_cells[2].text = 'Result'
    
    test_results = [
        ('Risk Assessment', 5, 'Effective'),
        ('Logical Access Control', 8, 'Effective'),
        ('System Monitoring', 4, 'Effective'),
        ('Disaster Recovery', 3, 'Effective'),
        ('Data Validation', 6, 'Effective'),
        ('Encryption Controls', 5, 'Effective'),
        ('Privacy Controls', 4, 'Effective'),
    ]
    
    for i, (objective, tests, result) in enumerate(test_results, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = objective
        row_cells[1].text = str(tests)
        row_cells[2].text = result
    
    doc.add_page_break()
    
    # Page 10: Conclusion and Management Responsibilities
    add_heading(doc, "Conclusion and Management Responsibilities", level=1)
    
    doc.add_paragraph(
        "In conclusion, Acme Cloud Services, Inc. has implemented and maintained effective controls related to Security, Availability, "
        "Processing Integrity, Confidentiality, and Privacy during the period January 1, 2023 to December 31, 2023."
    )
    
    doc.add_paragraph(
        "Management is responsible for: (1) selecting the appropriate Trust Service Criteria; (2) designing, implementing, and documenting controls; "
        "(3) ensuring controls are operating effectively; and (4) evaluating the effectiveness of controls and communicating results to authorized parties."
    )
    
    doc.add_paragraph(
        "We believe this report provides reasonable assurance regarding the operating effectiveness of the Company's controls during the specified period. "
        "However, we emphasize that controls may become less effective over time and that the occurrences of errors and irregularities, though unlikely, cannot be ruled out."
    )
    
    # Save the document
    output_path = "sample_soc2_report.docx"
    doc.save(output_path)
    print(f"✓ Fake SOC 2 report created: {output_path}")

if __name__ == "__main__":
    create_fake_soc2_report()
